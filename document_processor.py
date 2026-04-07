from pathlib import Path

from langchain_community.document_loaders import PDFPlumberLoader, PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import config


def _to_int(value, default: int):
    """Chuyển giá trị về int an toàn, dùng default nếu không hợp lệ."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return default

def _load_pdf(file_path: str):
    """Đọc PDF, ưu tiên PDFPlumber, dự phòng bằng PyPDF."""
    try:
        docs = PDFPlumberLoader(file_path).load()
    except Exception:
        docs = PyPDFLoader(file_path).load()

    if not docs:
        raise ValueError("Không đọc được nội dung PDF hoặc file rỗng.")

    # Chuẩn hóa metadata để mọi chunk đều có source + page nhất quán.
    source_name = Path(file_path).name
    raw_pages = []
    for doc in docs:
        page_candidate = (doc.metadata or {}).get("page")
        if page_candidate is None:
            page_candidate = (doc.metadata or {}).get("page_number")
        if page_candidate is not None:
            raw_pages.append(_to_int(page_candidate, 0))

    # Nhiều loader trả page bắt đầu từ 0, ta chuyển về 1-based để hiển thị cho người dùng.
    zero_based = bool(raw_pages) and min(raw_pages) == 0

    normalized_docs = []
    for index, doc in enumerate(docs, start=1):
        metadata = dict(doc.metadata or {})
        page_value = metadata.get("page", metadata.get("page_number", index))
        page_number = _to_int(page_value, index)
        if zero_based:
            page_number += 1

        metadata["page"] = max(1, page_number)
        metadata["source"] = source_name
        normalized_docs.append(Document(page_content=doc.page_content, metadata=metadata))

    return normalized_docs


def _load_docx(file_path: str):
    """Đọc DOCX bằng python-docx, giữ lại đoạn văn và nội dung bảng."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError(
            "Thiếu thư viện python-docx. Vui lòng cài: pip install python-docx"
        ) from exc

    doc = DocxDocument(file_path)
    sections = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                sections.append(" | ".join(cells))

    full_text = "\n".join(sections).strip()
    if not full_text:
        raise ValueError("Không đọc được nội dung DOCX hoặc file rỗng.")

    # DOCX không có khái niệm trang ổn định trong parser này, quy ước page=1 để tracking nguồn.
    return [
        Document(
            page_content=full_text,
            metadata={"source": Path(file_path).name, "type": "docx", "page": 1},
        )
    ]


def load_and_split_document(file_path: str, chunk_size: int | None = None, chunk_overlap: int | None = None):
    """Đọc tài liệu (PDF/DOCX) và chia nhỏ theo chunk strategy."""
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        docs = _load_pdf(file_path)
    elif suffix == ".docx":
        docs = _load_docx(file_path)
    else:
        raise ValueError("Định dạng file chưa được hỗ trợ. Vui lòng dùng PDF hoặc DOCX.")

    # Cắt nhỏ văn bản
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or config.CHUNK_SIZE,
        chunk_overlap=chunk_overlap or config.CHUNK_OVERLAP
    )
    chunks = text_splitter.split_documents(docs)

    # Bước khóa để bảo toàn metadata sau split: mỗi chunk bắt buộc có page + source.
    source_name = Path(file_path).name
    for chunk in chunks:
        metadata = chunk.metadata or {}
        metadata["source"] = metadata.get("source") or source_name
        metadata["page"] = _to_int(metadata.get("page", 1), 1)
        chunk.metadata = metadata

    return chunks


def load_and_split_pdf(file_path: str):
    """Giữ tương thích ngược với code cũ."""
    return load_and_split_document(file_path)